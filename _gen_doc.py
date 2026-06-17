from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime, os

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def code(doc, text, fs=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(fs)

def tb(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.cell(i, j).text = str(cell)
    return t

H = lambda d, t, l=1: d.add_heading(t, level=l)

# ---- 封面 ----
t = doc.add_heading('3D-PMGNet LiTS 肝脏肿瘤分割', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('代码结构与模块详解'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(f'生成: {datetime.date.today()}'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- 1 ----
H(doc, '1. 项目总览', 1)
doc.add_paragraph('基于 3D-PMGNet (Probability Map-Guided Network)，两阶段分割 LiTS 肝脏肿瘤：')
doc.add_paragraph('阶段一 (liver): CT → 肝脏区域 mask（肝脏+肿瘤合并为前景）', style='List Bullet')
doc.add_paragraph('阶段二 (tumor): [CT + 肝脏mask] → 肿瘤 mask（仅在肝脏区域内预测）', style='List Bullet')
tb(doc, [
    ['文件', '功能'],
    ['dataset_cadic.py', '数据加载: 重采样、HU窗口、ROI裁剪、两阶段标签'],
    ['MC_network_backbone.py', 'PMGNet 主网络: 编解码器 + 概率模块调度'],
    ['pmg_encoder.py', 'ConvNeXt 3D 骨干 (uxnet_conv + ux_block)'],
    ['PosFuse.py', '概率提示融合: 位置编码 + 空间门控注意力'],
    ['mc_refine.py', 'MC不确定性 + 动态阈值 + 高斯平滑 + 温度缩放'],
    ['train_two_stage.py', '训练: 断点/早停/Loss曲线'],
    ['predict.py', '推理: 滑动窗口 → LiTS提交格式'],
])

# ---- 2 ----
H(doc, '2. 数据预处理 (dataset_cadic.py)', 1)
doc.add_paragraph('预处理 5 步，将原始 CT → 固定尺寸张量：')

H(doc, '2.1 各向同性重采样', 2)
doc.add_paragraph('加载 NIfTI（SimpleITK，保留 spacing）→ 重采样到 (1.0,1.0,1.0)mm³。CT 用线性插值，标注用最近邻。')
code(doc, '''resampler = sitk.ResampleImageFilter()
resampler.SetOutputSpacing((1.0, 1.0, 1.0))
resampler.SetInterpolator(sitk.sitkLinear)          # CT
resampler.SetInterpolator(sitk.sitkNearestNeighbor) # 标注''')

H(doc, '2.2 HU 窗口 + 归一化', 2)
doc.add_paragraph('HU 窗口 [-160, 240] 肝脏软组织窗 → clip → 归一化到 [0,1]')
code(doc, 'img = np.clip(img, -160, 240)\nimg = (img - (-160)) / (240 - (-160))  # → [0,1]')

H(doc, '2.3 ROI 裁剪 + 标签处理', 2)
doc.add_paragraph('行 [20:428], 列 [92:418]，去除边缘无效区域。根据 stage 处理标签：')
tb(doc, [
    ['阶段', '标签', '输入通道'],
    ['liver', '原始 {0,1,2} → {0,1} (肝脏+肿瘤=1)', '1ch CT'],
    ['tumor', '原始 {2}→{1} 其余→0', '2ch [CT, liver_mask]'],
])

H(doc, '2.4 中心裁剪/Pad', 2)
doc.add_paragraph('目标尺寸默认 (1,96,96,96)。超过中心裁剪，不足对称补零。96³ 比 128³ 省约 42% 显存。')

# ---- 3 ----
H(doc, '3. 模型架构 (MC_network_backbone.py)', 1)
doc.add_paragraph('PMGNet: 编码器-解码器 + 概率图引导模块。37M 参数。')
code(doc, '''输入 (B, in_chans, 96, 96, 96)
  │
  ├─ uxnet_conv: 4级下采样 → [48,96,192,384]ch
  │     96³→48³→24³→12³→6³
  │
  ├─ Encoder ×5 (UnetrBasicBlock):
  │    enc1: in→48 (96³) ... enc5: 384→768 (6³, bottleneck)
  │
  ├─ MC + Refine (每层):
  │    5次MC Dropout→Softmax→平均→Refine精调→PosFuse融合
  │
  ├─ Decoder (Fuseblock ×4):
  │    转置卷积上采样 + Skip + PosFuse
  │    6³→12³→24³→48³→96³
  │
  └─ UnetOutBlock → (B, out_chans, 96, 96, 96)''')

# ---- 4 ----
H(doc, '4. 骨干网络 (pmg_encoder.py)', 1)
H(doc, '4.1 ux_block — ConvNeXt 基础块', 2)
doc.add_paragraph('Depthwise Conv(7³) → LayerNorm → 1×1 Conv(扩4倍) → GELU → 1×1 Conv(还原) + Layer Scale + 残差连接。')
code(doc, '''x = self.dwconv(x)              # 7³ depthwise (groups=dim)
x = x.permute(0,2,3,4,1)        # channels_last
x = self.norm(x)                 # LayerNorm
x = x.permute(0,4,1,2,3)        # channels_first
x = self.pwconv1(x)             # dim→4*dim (1³)
x = self.gelu(x)
x = self.pwconv2(x)             # 4*dim→dim (1³)
x = self.gamma * x + input      # Layer Scale + skip''')

H(doc, '4.2 uxnet_conv — 多级特征提取', 2)
doc.add_paragraph('stem Conv(7³, stride=2) → 4 stage，每 stage: downsample(stride=2) → depths[i]×ux_block → LayerNorm')
doc.add_paragraph('输出 4 尺度特征: [1/4, 1/8, 1/16, 1/32] 原始尺寸。depths=[2,2,2,2], dims=[48,96,192,384]。')

# ---- 5 ----
H(doc, '5. 概率图模块 — PMG 核心创新', 1)
doc.add_paragraph('论文核心贡献：用中间特征概率图引导分割过程。两个模块：')

H(doc, '5.1 MC + Refine (mc_refine.py)', 2)
doc.add_paragraph('对 encoder 特征做 MC Dropout → Refine 精炼 → 温度缩放。包含 4 个子步骤：')

doc.add_paragraph('① MC Dropout: 5次 softmax(dropout(feat)) → 平均 → 概率图。模拟贝叶斯推理，均值=预测，方差=不确定性。', style='List Number')
doc.add_paragraph('② ParamPredictor: 对每通道计算 (mean, var) → MLP → α(阈值) + σ(平滑宽度)。自适应学习。', style='List Number')
doc.add_paragraph('③ 动态阈值+高斯平滑: α分位数→阈值→高于阈值区域高斯平滑→混合。高置信度区域去噪，低置信度保留。', style='List Number')
doc.add_paragraph('④ 温度缩放: 每通道可学习 t_c → logits/t_c → sigmoid。控制概率"锐度"。', style='List Number')

code(doc, '''# 完整精炼流程
class RefineSegmentationMultiChannel(nn.Module):
    def forward(self, prob_map):          # (1,C,H,W,Z)
        for c in range(C):                # 逐通道处理
            mean, var = prob_single.mean(), prob_single.var()
            alpha, sigma = self.param_predictor(mean, var)
            refined = refine_segmentation(prob_single, alpha, sigma)
        refined = self.temp_scaling(refined)  # 温度缩放
        return refined''')

H(doc, '5.2 PosFuse 概率提示融合 (PosFuse.py)', 2)
doc.add_paragraph('将概率图作为"提示"，通过交叉注意力增强原始特征。包含 3 个子模块：')

doc.add_paragraph('① PositionEmbeddingRandom3D: 随机傅里叶特征生成 3D 位置编码。sin/cos(随机高斯投影) → 高维正交向量。', style='List Number')
doc.add_paragraph('② PromptEncoder: 2×Conv3d 编码概率图 → 加位置编码 → "提示特征"。输入 in_ch → 输出 in_ch×2。', style='List Number')
doc.add_paragraph('③ SpatialGate: cat([特征, 提示]) → Conv3d → Sigmoid → 空间注意力权重。公式: out = A ⊙ (1+σ(Conv([A,E_B])))。', style='List Number')

code(doc, '''class ProbPromptFusion(nn.Module):
    def forward(self, A, B):           # A: 特征 B: 概率图
        E_B = self.encoder(B)          # 编码概率图→提示特征
        attn = self.gate(E_B, A)       # 空间注意力权重
        return A * (1 + attn)          # 增强特征''')

doc.add_paragraph('所有 PosFuse 调用包裹 torch.utils.checkpoint (梯度检查点)，反向时重算而非存储中间激活，零精度损失省显存。')

# ---- 6 ----
H(doc, '6. 训练 (train_two_stage.py)', 1)
doc.add_paragraph('Loss = 0.5×DiceLoss + 0.5×CrossEntropyLoss')
tb(doc, [
    ['参数', '值'],
    ['epochs', '300'],
    ['lr', '1e-4 (AdamW)'],
    ['batch_size', '1 (MC Refine限制)'],
    ['AMP', 'GradScaler 混合精度'],
    ['patience', '50 (早停)'],
])
doc.add_paragraph('每个epoch: 自动保存 checkpoint.pth(模型+优化器+epoch) + history.json。Ctrl+C 不丢数据。')
doc.add_paragraph('训练结束生成 loss_curve.png (Train Loss + Val Dice 双栏图)。')

# ---- 7 ----
H(doc, '7. 推理 (predict.py)', 1)
doc.add_paragraph('滑动窗口 (96³, 50%重叠, 高斯权重融合) + 两阶段：')
doc.add_paragraph('加载CT→重采样→HU→ROI→Pad', style='List Number')
doc.add_paragraph('阶段一: liver_model→肝脏mask', style='List Number')
doc.add_paragraph('阶段二: [CT+liver_mask]→肿瘤mask', style='List Number')
doc.add_paragraph('合成标签(0/1/2)→逆向还原→保存NIfTI', style='List Number
