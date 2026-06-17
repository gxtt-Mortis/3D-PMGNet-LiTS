from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

doc = Document()
style = doc.styles["Normal"]
style.font.size = Pt(10.5)

def code(d, text, fs=8):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(fs)

def tb(d, rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.cell(i, j).text = str(cell)

H = lambda d, t, l=1: d.add_heading(t, level=l)

t = doc.add_heading("3D-PMGNet LiTS Code Documentation", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(f"Generated: {datetime.date.today()}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

H(doc, "1. Project Overview", 1)
doc.add_paragraph("Two-stage LiTS liver tumor segmentation using 3D-PMGNet.")
tb(doc, [
    ["File", "Purpose"],
    ["dataset_cadic.py", "Data: resample, HU window, ROI crop, two-stage labels"],
    ["MC_network_backbone.py", "PMGNet: encoder-decoder + probability modules"],
    ["pmg_encoder.py", "ConvNeXt 3D backbone (uxnet_conv + ux_block)"],
    ["PosFuse.py", "Probability prompt fusion: position encoding + spatial gate"],
    ["mc_refine.py", "MC uncertainty + dynamic threshold + Gaussian + temperature"],
    ["train_two_stage.py", "Training: checkpoint, early-stop, loss curves"],
])

H(doc, "2. Preprocessing (dataset_cadic.py)", 1)
doc.add_paragraph("5 steps: raw CT -> fixed tensor")
doc.add_paragraph("1. Resample to (1.0,1.0,1.0)mm isotropic (CT: linear, label: nearest)", style="List Number")
doc.add_paragraph("2. HU window [-160,240] -> clip -> normalize [0,1]", style="List Number")
doc.add_paragraph("3. ROI crop [20:428, 92:418] (optional)", style="List Number")
doc.add_paragraph("4. Stage label: liver={0,1}, tumor={0,1} from orig {0,1,2}", style="List Number")
doc.add_paragraph("5. Center crop/pad to (1,96,96,96). 96^3 saves ~42% VRAM vs 128^3", style="List Number")

H(doc, "3. PMGNet Architecture (MC_network_backbone.py)", 1)
doc.add_paragraph("37M params. Encoder-decoder + probability-guided modules.")
code(doc, """Input (B, in_chans, 96, 96, 96)
  |
  +-- uxnet_conv: 4-level -> [48,96,192,384]ch  (96^3->48^3->24^3->12^3->6^3)
  +-- Encoder x5: enc1(in->48,96^3) ... enc5(384->768,6^3)
  +-- MC+Refine per level: 5x Dropout->Softmax->mean->Refine->PosFuse
  +-- Decoder x4: TransposedConv + Skip + PosFuse (6^3->...->96^3)
  +-- Output: UnetOutBlock -> (B, out_chans, 96, 96, 96)""")

H(doc, "4. Backbone (pmg_encoder.py)", 1)
doc.add_paragraph("ux_block: ConvNeXt basic block. DWConv(7^3) -> LN -> 1x1(expand 4x) -> GELU -> 1x1(restore) + LayerScale.")
doc.add_paragraph("uxnet_conv: stem Conv(7^3,stride=2) -> 4 stages [stride=2 + depths[i]*ux_block + LN]. dims=[48,96,192,384].")

H(doc, "5. Probability Map Modules - Core Innovation", 1)
H(doc, "5.1 MC + Refine (mc_refine.py)", 2)
doc.add_paragraph("4 sub-steps making the probability-guided mechanism:")
doc.add_paragraph("1) MC Dropout: 5x softmax(dropout(feat)) -> mean. Bayesian-style.", style="List Number")
doc.add_paragraph("2) ParamPredictor: per-channel (mean,var) -> MLP -> alpha(threshold) + sigma(smooth).", style="List Number")
doc.add_paragraph("3) Dynamic threshold + Gaussian smooth: threshold by alpha-quantile, smooth high-conf regions.", style="List Number")
doc.add_paragraph("4) Temperature scaling: per-channel learnable t_c -> sharpen/soften distribution.", style="List Number")

H(doc, "5.2 PosFuse (PosFuse.py)", 2)
doc.add_paragraph("Probability map as prompt to enhance features via cross-attention. 3 sub-modules:")
doc.add_paragraph("1) PositionEmbeddingRandom3D: random Fourier features for 3D position encoding.", style="List Number")
doc.add_paragraph("2) PromptEncoder: 2xConv3d + position encoding. in_ch -> in_ch*2 prompt features.", style="List Number")
doc.add_paragraph("3) SpatialGate: cat([feature,prompt]) -> Conv3d -> Sigmoid. out = A*(1+sigma(Conv([A,E_B]))).", style="List Number")
doc.add_paragraph("Gradient checkpointing applied: backward recomputes instead of storing. Zero accuracy loss.")

H(doc, "6. Training (train_two_stage.py)", 1)
doc.add_paragraph("Loss = 0.5*Dice + 0.5*CE. AdamW(1e-4), AMP GradScaler, patience=50 early stop.")
doc.add_paragraph("Auto-save: checkpoint.pth (model+optimizer+epoch) every epoch. Ctrl+C safe.")
doc.add_paragraph("Post-training: history.json + loss_curve.png generated.")

H(doc, "7. Inference (predict.py)", 1)
doc.add_paragraph("Sliding window (96^3, 50% overlap, Gaussian weight) + two-stage: CT->liver->tumor->save NIfTI.")

H(doc, "8. Commands", 1)
code(doc, """cd PMGNet
python train_two_stage.py --stage all --epochs 300 --batch_size 1 --patience 50
python train_two_stage.py --stage liver --resume --batch_size 1
python predict.py --input_dir ../test_data --output_dir ../predictions""")

save = r"E:\肝脏肿瘤分割\3D-PMGNet-main\3D-PMGNet_代码详解.docx"
doc.save(save)
print(f"OK: {os.path.getsize(save):,} bytes")
