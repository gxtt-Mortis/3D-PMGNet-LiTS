import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '等线'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0, 70, 127)
    elif level == 3:
        heading_style.font.size = Pt(12)
        heading_style.font.color.rgb = RGBColor(0, 90, 150)


def add_code_block(doc, code_text, language="Python"):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)

    # 添加灰色背景的代码框
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
#                         封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('3D-PMGNet 项目代码解析文档')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Probability Map-Guided Network for 3D\nVolumetric Medical Image Segmentation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('论文来源：IEEE Transactions on Image Processing, 2025\n生成日期：2026年6月15日')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
#                         目录页
# ============================================================
doc.add_heading('目  录', level=1)
doc.add_paragraph()

toc_items = [
    ('一、', '项目概述与论文背景'),
    ('二、', '项目文件结构总览'),
    ('三、', '环境配置与依赖'),
    ('四、', '核心模块详解'),
    ('', '4.1  pmg_encoder.py — 3D 特征编码器（ConvNeXt骨干网络）'),
    ('', '4.2  MC_network_backbone.py — 主网络架构（PMGNet）'),
    ('', '4.3  mc_refine.py — 概率图精炼模块'),
    ('', '4.4  PosFuse.py — 位置感知概率融合模块'),
    ('五、', '数据处理与训练'),
    ('', '5.1  dataset_cadic.py — 数据加载器'),
    ('', '5.2  train_cadic.py — 训练脚本'),
    ('六、', '整体数据流与运行流程'),
    ('七、', '核心创新点总结'),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f'{num}{item}')
        run.font.bold = True
    else:
        run = p.add_run(f'     {item}')
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
#              一、项目概述与论文背景
# ============================================================
doc.add_heading('一、项目概述与论文背景', level=1)

doc.add_paragraph(
    '3D-PMGNet（Probability Map-Guided Network for 3D Volumetric Medical Image Segmentation）'
    '是一个用于3D医学图像分割的深度学习模型，发表于 IEEE Transactions on Image Processing (2025)。'
    '该模型专门针对CT等3D医学影像中肝脏肿瘤等病灶的自动分割任务设计。'
)

doc.add_heading('论文核心动机', level=3)
doc.add_paragraph(
    '3D医学图像（如CT扫描）具有以下固有挑战：'
)

problems = [
    '各向异性（Anisotropy）：不同方向的采样分辨率不一致，导致某些方向上图像细节模糊或失真',
    '强度不均匀性（Intensity Inhomogeneity）：由于设备限制、扫描参数或患者解剖差异，图像中出现全局或局部的亮度不均匀',
    '边界模糊：上述问题导致病灶边界不清晰，模型容易聚焦到不相关区域',
]
for i, prob in enumerate(problems, 1):
    doc.add_paragraph(f'{i}. {prob}', style='List Number')

doc.add_paragraph(
    '为应对这些挑战，3D-PMGNet 提出了"概率图引导"（Probability Map Guidance）的核心思想：'
    '利用网络中间层生成的概率图作为自监督信号，引导分割过程聚焦于可靠的高响应区域，同时抑制低响应噪声。'
)

doc.add_heading('论文引用', level=3)
add_code_block(doc,
    '@article{zhuprobability,\n'
    '  title={Probability Map-Guided Network for 3D Volumetric Medical Image Segmentation},\n'
    '  author={Zhu, Zhiqin and Zhang, Zimeng and Qi, Guanqiu and Li, Yuanyuan and Yang, Pan and Liu, Yu},\n'
    '  journal={IEEE Transactions on Image Processing},\n'
    '  year={2025},\n'
    '  volume={34},\n'
    '  pages={7222-7234},\n'
    '  doi={10.1109/TIP.2025.3623259},\n'
    '  pmid={41187034}\n'
    '}'
)

doc.add_page_break()

# ============================================================
#              二、项目文件结构总览
# ============================================================
doc.add_heading('二、项目文件结构总览', level=1)

doc.add_paragraph('以下是项目的完整文件结构及各文件功能说明：')

add_code_block(doc,
    '3D-PMGNet-main/\n'
    '├── README.md                                    # 项目说明与论文引用\n'
    '└── PMGNet/\n'
    '    ├── dataset_cadic.py                         # 数据集加载与预处理\n'
    '    ├── train_cadic.py                           # 训练主脚本\n'
    '    ├── runs/ct/                                 # TensorBoard 训练日志\n'
    '    │   └── events.out.tfevents.*\n'
    '    └── PMGNet/                                  # 核心模型代码\n'
    '        ├── __init__.py                          # (缺失，建议添加)\n'
    '        ├── pmg_encoder.py                       # 3D ConvNeXt 编码器\n'
    '        ├── MC_network_backbone.py               # PMGNet 主网络\n'
    '        ├── mc_refine.py                         # 概率图精炼模块\n'
    '        └── PosFuse.py                           # 位置感知融合模块'
)

add_table_with_style(doc,
    ['文件', '类别', '功能描述'],
    [
        ['pmg_encoder.py', '模型核心', '3D ConvNeXt 骨干网络，负责多尺度特征提取'],
        ['MC_network_backbone.py', '模型核心', 'PMGNet 主网络，U-Net风格编码器-解码器'],
        ['mc_refine.py', '模型核心', '蒙特卡洛采样 + 动态阈值 + 高斯平滑优化概率图'],
        ['PosFuse.py', '模型核心', '位置编码 + Prompt编码 + 空间门控特征融合'],
        ['train_cadic.py', '训练', '训练流程管理、损失计算、指标评估、模型保存'],
        ['dataset_cadic.py', '数据处理', 'NIfTI格式数据加载、裁剪/填充预处理'],
    ]
)

doc.add_page_break()

# ============================================================
#              三、环境配置与依赖
# ============================================================
doc.add_heading('三、环境配置与依赖', level=1)

doc.add_heading('3.1 当前已安装的核心包', level=2)

add_table_with_style(doc,
    ['包名', '版本', '用途'],
    [
        ['torch', '2.12.0', '深度学习框架（PyTorch）'],
        ['torchvision', '0.27.0', 'PyTorch 视觉工具库'],
        ['torchaudio', '2.11.0', 'PyTorch 音频工具库'],
        ['numpy', '2.4.6', '科学计算库'],
        ['nibabel', '5.4.2', 'NIfTI 格式医学图像读写'],
        ['SimpleITK', '2.5.5', '医学图像处理工具包'],
        ['timm', '1.0.22', 'PyTorch 图像模型库（提供 DropPath 等）'],
        ['tqdm', '4.68.2', '进度条显示'],
    ]
)

doc.add_heading('3.2 仍需安装的包', level=2)

doc.add_paragraph('以下包在当前环境中尚未安装，运行项目前需要安装：')

add_table_with_style(doc,
    ['包名', '预计版本', '用途', '安装命令'],
    [
        ['monai', '≥1.3.0', '医学AI框架（UNETR模块、Dice Loss等）', 'pip install monai'],
        ['tensorboard', '≥2.14.0', '训练可视化（SummaryWriter）', 'pip install tensorboard'],
        ['python-docx', '≥1.0.0', 'Word文档生成（仅本文档使用）', 'pip install python-docx'],
    ]
)

doc.add_heading('3.3 完整环境配置步骤', level=2)

doc.add_paragraph('推荐使用 Conda 创建独立环境，按以下步骤配置：')

steps = [
    '# 步骤1：创建 Conda 环境（Python 3.9+）',
    'conda create -n pmgnet python=3.9 -y',
    'conda activate pmgnet',
    '',
    '# 步骤2：安装 PyTorch（根据CUDA版本选择）',
    '# CUDA 11.8 版本：',
    'pip install torch==2.1.0 torchvision==0.16.0 --index-url https://download.pytorch.org/whl/cu118',
    '# CUDA 12.1 版本：',
    'pip install torch==2.1.0 torchvision==0.16.0 --index-url https://download.pytorch.org/whl/cu121',
    '',
    '# 步骤3：安装医学图像处理库',
    'pip install monai>=1.3.0',
    'pip install SimpleITK>=2.3.0',
    'pip install nibabel>=5.0.0',
    '',
    '# 步骤4：安装其他依赖',
    'pip install timm>=0.9.0',
    'pip install tensorboard>=2.14.0',
    'pip install tqdm>=4.65.0',
    'pip install numpy>=1.24.0',
]
add_code_block(doc, '\n'.join(steps))

doc.add_heading('3.4 requirements.txt（推荐创建）', level=2)

doc.add_paragraph('建议在项目根目录创建 requirements.txt 文件，内容如下：')

add_code_block(doc,
    'torch>=2.0.0\n'
    'torchvision>=0.15.0\n'
    'monai>=1.3.0\n'
    'SimpleITK>=2.3.0\n'
    'nibabel>=5.0.0\n'
    'timm>=0.9.0\n'
    'tensorboard>=2.14.0\n'
    'tqdm>=4.65.0\n'
    'numpy>=1.24.0'
)

doc.add_page_break()

# ============================================================
#              四、核心模块详解
# ============================================================
doc.add_heading('四、核心模块详解', level=1)

# ---------- 4.1 pmg_encoder.py ----------
doc.add_heading('4.1  pmg_encoder.py — 3D特征编码器（ConvNeXt骨干网络）', level=2)

doc.add_heading('功能定位', level=3)
doc.add_paragraph(
    '该文件实现了基于 ConvNeXt 架构的3D特征提取骨干网络，是整个模型的特征提取基础。'
    '输入3D医学图像，输出4个不同分辨率的多尺度特征图。'
)

doc.add_heading('类与函数说明', level=3)

add_table_with_style(doc,
    ['类/函数', '行号', '功能'],
    [
        ['LayerNorm', '7-33', '支持 channels_last 和 channels_first 两种数据格式的层归一化。channels_last 用于 (B,H,W,D,C) 格式，channels_first 用于 (B,C,H,W,D) 格式'],
        ['ux_block', '35-75', 'ConvNeXt 基础模块：深度可分离卷积(7×7×7) → 层归一化 → 1×1逐点卷积(GELU) → 1×1逐点卷积 → 残差连接 + DropPath'],
        ['uxnet_conv', '78-149', '完整的3D ConvNeXt 骨干网络，包含4个阶段的 stem+下采样+ux_block堆叠'],
    ]
)

doc.add_heading('关键参数', level=3)

add_table_with_style(doc,
    ['参数', '默认值', '说明'],
    [
        ['in_chans', '1', '输入通道数（CT灰度图为1）'],
        ['depths', '[2,2,2,2]', '4个阶段各自的 ux_block 数量'],
        ['dims', '[48,96,192,384]', '4个阶段的特征通道数'],
        ['drop_path_rate', '0.0', 'Stochastic Depth 比率'],
        ['out_indices', '[0,1,2,3]', '输出哪些阶段的特征图'],
    ]
)

doc.add_heading('数据流', level=3)
doc.add_paragraph(
    '输入: (B, 1, D, H, W) → Stem(Conv7×7, stride=2) → Stage0(2个ux_block) → Norm0 → out[0] (B,48,D/2,H/2,W/2)\n'
    '→ Downsample1 → Stage1(2个ux_block) → Norm1 → out[1] (B,96,D/4,H/4,W/4)\n'
    '→ Downsample2 → Stage2(2个ux_block) → Norm2 → out[2] (B,192,D/8,H/8,W/8)\n'
    '→ Downsample3 → Stage3(2个ux_block) → Norm3 → out[3] (B,384,D/16,H/16,W/16)\n\n'
    '每个 Downsample 层为: LayerNorm + Conv3d(kernel=2, stride=2)，将空间尺寸减半、通道数翻倍。'
)

doc.add_heading('ux_block 详解', level=3)
doc.add_paragraph(
    'ux_block 是 ConvNeXt 的核心构建块，其设计灵感来自 ConvNeXt 论文。关键设计特点：'
)

features = [
    '深度可分离卷积（Depthwise Conv）：使用 groups=dim 的 7×7×7 3D卷积，大幅减少参数量',
    '倒置瓶颈结构：中间层扩展为 4×dim 通道，使用 GELU 激活函数',
    'Layer Scale：可学习的缩放参数 gamma，初始值极小(1e-6)，稳定训练',
    'Stochastic Depth（DropPath）：训练时随机丢弃整个 block，作为正则化手段',
    '通道优先策略：使用 channels_last 格式进行 LayerNorm（性能更优），然后转回 channels_first 进行卷积',
]
for f in features:
    doc.add_paragraph(f'• {f}')

doc.add_page_break()

# ---------- 4.2 MC_network_backbone.py ----------
doc.add_heading('4.2  MC_network_backbone.py — 主网络架构（PMGNet）', level=2)

doc.add_heading('功能定位', level=3)
doc.add_paragraph(
    '该文件是项目的核心文件，定义了完整的 PMGNet 分割网络。'
    '采用 U-Net 风格的编码器-解码器架构，集成了概率图引导机制。'
    '包含两个主要类：Fuseblock（特征融合模块）和 PMGNet（主网络）。'
)

doc.add_heading('Fuseblock 类', level=3)

add_table_with_style(doc,
    ['组件', '功能'],
    [
        ['转置卷积 (transp_conv)', '对输入特征进行2倍上采样（kernel_size=2, stride=2）'],
        ['空间对齐 (interpolate)', '使用三线性插值将跳跃连接特征对齐到上采样后的尺寸'],
        ['ProbPromptFusion', '通过概率提示编码和空间门控机制融合特征'],
    ]
)

doc.add_paragraph(
    'Fuseblock 前向流程：\n'
    '1. 转置卷积上采样高层特征 (inp)\n'
    '2. 三线性插值对齐跳跃连接特征 (skip) 的尺寸\n'
    '3. 调用 ProbPromptFusion 进行概率引导的特征融合'
)

doc.add_heading('PMGNet 类 — 网络架构', level=3)

doc.add_paragraph('PMGNet 由以下组件构成：')

components = [
    ('编码器（Encoder）',
     'uxnet_conv: 3D ConvNeXt 骨干网络，提取4个尺度的特征\n'
     'encoder1~5: 5个 UNETR BasicBlock，逐步变换特征（通道数: 1→48→96→192→384→768）'),
    ('概率图精炼（MC Refine）',
     'refine1~4: 4个 RefineSegmentationMultiChannel 模块\n'
     '对每个编码器输出的中间特征，执行蒙特卡洛采样 → softmax → 平均 → 动态阈值+高斯平滑'),
    ('概率融合（Prob Fusion）',
     'ProbPromptFusion: 将精炼后的概率图编码为提示（Prompt），通过3D位置编码增强后，与主特征进行空间门控融合'),
    ('解码器（Decoder）',
     'decoder5~1: 4个 Fuseblock + 1个 UNETR BasicBlock\n'
     '逐级上采样并融合跳跃连接特征（通道数: 768→384→192→96→48→48）'),
    ('输出层（Output）',
     'UnetOutBlock: 将最终特征映射到类别数（如4类），输出分割结果'),
]
for name, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.font.bold = True
    p.add_run(desc)

doc.add_heading('关键参数', level=3)

add_table_with_style(doc,
    ['参数', '默认值', '说明'],
    [
        ['in_chans', '1', '输入通道数'],
        ['out_chans', '13', '输出类别数'],
        ['depths', '[2,2,2,2]', '各阶段 ConvNeXt block 数量'],
        ['feat_size', '[48,96,192,384]', '各阶段特征通道数'],
        ['hidden_size', '768', '最深层的隐藏特征通道数'],
        ['spatial_dims', '3', '空间维度（3D）'],
        ['mc_samples', '5', '蒙特卡洛采样次数'],
        ['use_mc_refine', 'True', '是否启用 MC+Refine 概率图优化'],
    ]
)

doc.add_heading('前向传播流程', level=3)

steps = [
    '1. uxnet_3d(x_in) → 提取4个多尺度特征 outs[0..3]',
    '2. encoder1(x_in) → MC采样+精炼 → ProbPromptFusion → re_enc1 (跳跃连接)',
    '3. encoder2(outs[0]) → MC采样+精炼 → ProbPromptFusion → re_enc2 (跳跃连接)',
    '4. encoder3(outs[1]) → MC采样+精炼 → ProbPromptFusion → re_enc3 (跳跃连接)',
    '5. encoder4(outs[2]) → MC采样+精炼 → ProbPromptFusion → re_enc4 (跳跃连接)',
    '6. encoder5(outs[3]) → enc_hidden (最深特征, 768通道)',
    '7. decoder5(enc_hidden, re_enc4) → dec3 → decoder4(dec3, re_enc3) → ... → decoder1(dec0)',
    '8. out(out_feat) → 最终分割结果 (B, out_chans, D, H, W)',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ---------- 4.3 mc_refine.py ----------
doc.add_heading('4.3  mc_refine.py — 概率图精炼模块', level=2)

doc.add_heading('功能定位', level=3)
doc.add_paragraph(
    '该文件实现了概率图的自适应精炼（Refinement），是3D-PMGNet的核心创新之一。'
    '通过动态阈值、局部自适应高斯平滑和可学习温度缩放，提升概率图的可靠性。'
)

doc.add_heading('类与函数说明', level=3)

add_table_with_style(doc,
    ['类/函数', '行号', '功能'],
    [
        ['ParamPredictor', '5-19', '小型MLP网络，输入每个通道的均值和方差，预测动态阈值参数 alpha 和 高斯平滑参数 sigma'],
        ['dynamic_threshold()', '21-26', '根据 alpha 参数计算概率图的动态分位数阈值'],
        ['gaussian_kernel_3d()', '28-33', '生成3D高斯平滑核'],
        ['gaussian_weighted_smoothing()', '36-44', '对概率图进行3D高斯加权平滑'],
        ['refine_segmentation()', '46-51', '核心精炼函数：动态阈值分割 → 高斯平滑 → 掩码融合'],
        ['TemperatureScaling', '53-71', '可学习的通道级温度缩放，调整概率分布使其更接近真实分布'],
        ['RefineSegmentationMultiChannel', '74-105', '多通道概率图精炼主模块，整合上述所有组件'],
    ]
)

doc.add_heading('精炼流程详解', level=3)

doc.add_paragraph('对输入概率图 (1, C, H, W, D)，逐通道执行以下操作：')

steps = [
    '1. 计算通道均值和方差 → 输入 ParamPredictor MLP',
    '2. MLP输出 alpha (sigmoid, 0~1) 和 sigma (exp, >0)',
    '3. dynamic_threshold: 取概率图的 alpha 分位数作为阈值',
    '4. gaussian_weighted_smoothing: 用 sigma 控制的高斯核对概率图做平滑',
    '5. 动态掩码融合: refined = mask * smoothed + (1-mask) * original',
    '6. TemperatureScaling: 对精炼后的概率图应用通道级温度缩放',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('核心公式', level=3)
doc.add_paragraph(
    '精炼公式:   P_refined = M * S(P, σ) + (1-M) * P\n\n'
    '其中:\n'
    '  P = 原始概率图\n'
    '  M = 动态掩码: M(x) = 1 当 P(x) ≥ threshold(α), 否则 0\n'
    '  S(P, σ) = 以σ为参数的高斯平滑\n'
    '  α = 动态阈值参数 (由MLP预测)\n'
    '  σ = 高斯平滑标准差 (由MLP预测)\n\n'
    '温度缩放:   P_scaled = softmax(logit(P) / t_c)\n'
    '其中 t_c = exp(u_c) 是通道 c 的可学习温度参数'
)

doc.add_page_break()

# ---------- 4.4 PosFuse.py ----------
doc.add_heading('4.4  PosFuse.py — 位置感知概率融合模块', level=2)

doc.add_heading('功能定位', level=3)
doc.add_paragraph(
    '该文件实现了基于概率图的位置感知特征融合（ProbPromptFusion），是3D-PMGNet的另一个核心创新。'
    '将概率图编码为"提示"（Prompt），通过3D随机傅里叶位置编码增强空间信息，'
    '最后通过空间门控机制动态调整主特征的响应强度。'
)

doc.add_heading('类与函数说明', level=3)

add_table_with_style(doc,
    ['类', '行号', '功能'],
    [
        ['PositionEmbeddingRandom3D', '8-47', '3D随机傅里叶位置编码。使用随机高斯矩阵将坐标投影到高维空间，然后应用 sin/cos 编码'],
        ['PromptEncoder', '50-70', '将概率图编码为"提示"特征：Conv3d(3×3) → GroupNorm → GELU → Conv3d(1×1) → + 位置编码'],
        ['SpatialGate', '73-103', '空间门控模块：拼接主特征和提示特征 → Conv3d → Sigmoid → 得到注意力图'],
        ['ProbPromptFusion', '106-127', '主融合模块：Encoder编码概率图 → Gate生成注意力 → A * (1 + attention)'],
    ]
)

doc.add_heading('融合流程详解', level=3)

doc.add_paragraph(
    'ProbPromptFusion.forward(A, B) 的执行流程：\n\n'
    '输入:\n'
    '  A = 主特征图 (B, C_feat, D, H, W)  — 来自编码器的特征\n'
    '  B = 概率图 (B, C_prob, D, H, W)    — 来自MC精炼的概率图\n\n'
    '步骤:\n'
    '  1. PromptEncoder(B) → E_B (B, C_prob*2, D, H, W)\n'
    '     - Conv3d 提取特征\n'
    '     - 加上 3D 随机傅里叶位置编码\n\n'
    '  2. SpatialGate(E_B, A) → attention (B, 1, D, H, W)\n'
    '     - 拼接 [A, E_B] 后通过 Conv3d → Sigmoid\n'
    '     - 空间尺寸自动对齐\n\n'
    '  3. 输出: A * (1 + attention)\n'
    '     - attention 值在 0~1 之间\n'
    '     - 高响应区域被增强（×1~×2），低响应区域保持不变'
)

doc.add_heading('位置编码原理', level=3)
doc.add_paragraph(
    'PositionEmbeddingRandom3D 使用随机傅里叶特征（Random Fourier Features）进行位置编码：\n\n'
    '  PE(x) = [sin(2π·B·x), cos(2π·B·x)]\n\n'
    '其中:\n'
    '  x = 3D归一化坐标 [-1, 1]\n'
    '  B = 随机高斯矩阵 (3 × embed_dim/2)，scale = 1/√3\n'
    '  最终维度 = embed_dim\n\n'
    '这种编码方式能够让网络感知到3D空间中的绝对位置信息，'
    '有助于理解病灶在器官中的相对位置关系。'
)

doc.add_page_break()

# ============================================================
#              五、数据处理与训练
# ============================================================
doc.add_heading('五、数据处理与训练', level=1)

# ---------- 5.1 dataset_cadic.py ----------
doc.add_heading('5.1  dataset_cadic.py — 数据加载器', level=2)

doc.add_heading('功能概述', level=3)
doc.add_paragraph(
    'CTCACSimpleDataset 类继承自 PyTorch Dataset，负责加载和预处理3D医学图像数据。'
    '支持 NIfTI (.nii.gz) 格式，自动处理不同尺寸的图像。'
)

doc.add_heading('数据目录结构', level=3)
add_code_block(doc,
    'data_root/\n'
    '├── train/\n'
    '│   ├── imagesTr/\n'
    '│   │   ├── image_001.nii.gz\n'
    '│   │   └── image_002.nii.gz\n'
    '│   └── labelsTr/\n'
    '│       ├── label_001.nii.gz\n'
    '│       └── label_002.nii.gz\n'
    '└── val/\n'
    '    ├── imagesTr/\n'
    '    │   └── image_003.nii.gz\n'
    '    └── labelsTr/\n'
    '        └── label_003.nii.gz'
)

doc.add_heading('关键参数', level=3)

add_table_with_style(doc,
    ['参数', '默认值', '说明'],
    [
        ['data_root', '(必填)', '数据根目录路径'],
        ['phase', '"train"', '数据集阶段: train 或 val'],
        ['target_shape', '(1,128,128,128)', '统一的目标尺寸 (C,D,H,W)'],
        ['label_suffix', '"label"', '标签文件前缀，如 "labelnew"'],
    ]
)

doc.add_heading('预处理流程', level=3)
steps = [
    '1. 扫描 imagesTr 目录，提取所有 case_id',
    '2. 使用 SimpleITK（优先）或 NiBabel 加载 .nii.gz 文件',
    '3. 转换为 (C, D, H, W) 格式的 PyTorch Tensor',
    '4. 中心裁剪：如果尺寸大于目标，从中心裁掉多余部分',
    '5. 对称填充：如果尺寸小于目标，在四周对称填充0',
    '6. 返回 {"image": tensor, "label": tensor, "case_id": str} 字典',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ---------- 5.2 train_cadic.py ----------
doc.add_heading('5.2  train_cadic.py — 训练脚本', level=2)

doc.add_heading('功能概述', level=3)
doc.add_paragraph(
    'CTTrainer 类封装了完整的训练流程，包括模型构建、数据加载、损失计算、'
    '混合精度训练（AMP）、指标评估、TensorBoard可视化和模型保存。'
)

doc.add_heading('训练配置参数', level=3)

add_table_with_style(doc,
    ['参数', '默认值', '说明'],
    [
        ['data_root', '(必填)', '数据根目录'],
        ['num_classes', '4', '分割类别数'],
        ['epochs', '300', '训练总轮数'],
        ['lr', '1e-4', 'AdamW 初始学习率'],
        ['batch_size', '1 (代码默认)', '批次大小（__main__中设为4）'],
        ['alpha', '0.5', 'Dice Loss 权重（CE Loss 权重为 1-alpha）'],
        ['num_workers', '4', 'DataLoader 工作线程数'],
        ['target_shape', '(1,128,128,128)', '输入图像统一尺寸'],
    ]
)

doc.add_heading('损失函数', level=3)
doc.add_paragraph(
    '采用组合损失函数:\n\n'
    '  Total Loss = α × Dice Loss + (1-α) × CrossEntropy Loss\n\n'
    '  - Dice Loss (MONAI): 基于 Dice 系数的区域重叠损失，to_onehot_y=True, softmax=True, include_background=False\n'
    '  - CrossEntropy Loss: 像素级分类损失，ignore_index=-100\n'
    '  - α = 0.5: 两种损失等权重\n\n'
    '这种组合同时优化了区域重叠（Dice）和像素分类精度（CE），互补各自的不足。'
)

doc.add_heading('训练技巧', level=3)
features = [
    '混合精度训练（AMP）：使用 GradScaler + autocast，减少显存占用，加速训练',
    'AdamW 优化器：带权重衰减的 Adam，weight_decay=1e-5，防止过拟合',
    'cudnn.benchmark = True：自动寻找最优卷积算法',
    'pin_memory=True：加速 CPU→GPU 数据传输',
    '最佳模型保存：基于验证集平均 Dice 系数自动保存最优权重',
    'TensorBoard 可视化：记录训练损失、各类别 Dice 指标',
]
for f in features:
    doc.add_paragraph(f'• {f}')

doc.add_heading('评估指标', level=3)
doc.add_paragraph(
    '对每个类别 c（排除背景类0）计算 Dice 系数:\n\n'
    '  Dice_c = 2 × |Pred_c ∩ Target_c| / (|Pred_c| + |Target_c|)\n\n'
    '使用平均 Dice（各类别 Dice 的均值）作为模型选择标准。'
)

doc.add_page_break()

# ============================================================
#              六、整体数据流与运行流程
# ============================================================
doc.add_heading('六、整体数据流与运行流程', level=1)

doc.add_heading('6.1 完整数据流图', level=2)

doc.add_paragraph(
    '以下展示从输入到输出的完整数据流（假设输入 128×128×128 的CT图像，4类分割）：'
)

add_code_block(doc,
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                    输入: CT图像 (1, 1, 128, 128, 128)            │\n'
    '└───────────────────────────┬─────────────────────────────────────┘\n'
    '                            │\n'
    '    ┌───────────────────────▼──────────────────────────┐\n'
    '    │          uxnet_conv (3D ConvNeXt)                 │\n'
    '    │  Stem → Stage0→3, 每次下采样2×, 通道翻倍          │\n'
    '    └───┬──────────┬──────────┬──────────┬─────────────┘\n'
    '        │          │          │          │\n'
    '   out[0]     out[1]     out[2]     out[3]\n'
    '  (B,48,     (B,96,     (B,192,    (B,384,\n'
    '   64,64,64)  32,32,32)  16,16,16)  8,8,8)\n'
    '        │          │          │          │\n'
    '   encoder1    encoder2   encoder3   encoder4   encoder5\n'
    '  (1→48)     (48→96)   (96→192)  (192→384)  (384→768)\n'
    '        │          │          │          │\n'
    '   MC refine1  MC refine2 MC refine3 MC refine4\n'
    '  (5次采样)   (5次采样)  (5次采样)  (5次采样)\n'
    '        │          │          │          │\n'
    '   ProbFuse1   ProbFuse2  ProbFuse3  ProbFuse4\n'
    '        │          │          │          │\n'
    '   re_enc1    re_enc2    re_enc3    re_enc4   enc_hidden\n'
    '  (B,48,     (B,96,     (B,192,    (B,384,   (B,768,\n'
    '   64,64,64)  32,32,32)  16,16,16)  8,8,8)    4,4,4)\n'
    '        │          │          │          │          │\n'
    '        │          │          │     ┌────┴──────────┘\n'
    '        │          │          │     │ decoder5(Fuseblock)\n'
    '        │          │          │     │  上采样+融合re_enc4\n'
    '        │          │          │     └────────┬──────────┘\n'
    '        │          │          │          dec3 (B,192,8,8,8)\n'
    '        │          │     ┌────┴──────────┘\n'
    '        │          │     │ decoder4(Fuseblock)\n'
    '        │          │     │  上采样+融合re_enc3\n'
    '        │          │     └────────┬──────────┘\n'
    '        │          │          dec2 (B,96,16,16,16)\n'
    '        │     ┌────┴──────────┘\n'
    '        │     │ decoder3(Fuseblock)\n'
    '        │     │  上采样+融合re_enc2\n'
    '        │     └────────┬──────────┘\n'
    '        │          dec1 (B,48,32,32,32)\n'
    '   ┌────┴──────────┘\n'
    '   │ decoder2(Fuseblock)\n'
    '   │  上采样+融合re_enc1\n'
    '   └────────┬──────────┘\n'
    '        dec0 (B,48,64,64,64)\n'
    '            │\n'
    '       decoder1(UNETR BasicBlock)\n'
    '            │\n'
    '       out_feat (B,48,64,64,64)\n'
    '            │\n'
    '       UnetOutBlock\n'
    '            │\n'
    '   ┌────────▼──────────────────────────────────────────┐\n'
    '   │      输出: 分割结果 (1, 4, 64, 64, 64)             │\n'
    '   │      4个通道分别对应: 背景/肝脏/肿瘤/其他           │\n'
    '   └───────────────────────────────────────────────────┘'
)

doc.add_heading('6.2 训练流程', level=2)

steps = [
    '数据准备：CTCACSimpleDataset 加载 NIfTI 图像，统一尺寸到 128³',
    '模型初始化：PMGNet 构建完整网络，迁移到 GPU',
    '训练循环（300 epochs）：',
    '  a. 前向传播：图像 → PMGNet → 4通道分割输出',
    '  b. 损失计算：Dice Loss(0.5) + CE Loss(0.5)',
    '  c. 反向传播：AMP 混合精度 + GradScaler',
    '  d. 指标计算：各类别 Dice 系数',
    '验证循环：',
    '  a. 不计算梯度，纯推理模式',
    '  b. 计算验证集各类别 Dice 系数',
    '模型保存：',
    '  a. 平均 Dice 超过历史最佳 → 保存 best_ct_model.pth',
    'TensorBoard 记录：训练损失 + 训练/验证 Dice 指标',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('6.3 推理流程', level=2)

steps = [
    '1. 加载训练好的模型权重 best_ct_model.pth',
    '2. 加载待分割的 CT 图像，预处理到 128×128×128',
    '3. 模型前向传播（model.eval()，use_mc_refine=True）',
    '4. 对输出做 argmax 得到每个像素的类别标签',
    '5. 后处理（可选）：最大连通域分析、小目标去除等',
    '6. 将结果保存为 NIfTI 格式或可视化',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ============================================================
#              七、核心创新点总结
# ============================================================
doc.add_heading('七、核心创新点总结', level=1)

doc.add_paragraph(
    '3D-PMGNet 包含以下四个核心创新，分别对应不同的代码模块：'
)

innovations = [
    ('创新1：概率图引导机制 (Probability Map Guidance)',
     '文件: MC_network_backbone.py, mc_refine.py\n\n'
     '利用网络中间层特征生成概率图，将其作为自监督信号引导分割过程。'
     '与传统方法不同，概率图不是最终输出，而是作为中间"提示"反馈给网络，'
     '帮助网络聚焦于高置信度区域，同时抑制低响应噪声。\n\n'
     '实现方式: encoder输出 → MC采样 → 精炼 → ProbPromptFusion → 回馈给解码器'),

    ('创新2：自适应概率精炼 (Adaptive Probability Refinement)',
     '文件: mc_refine.py\n\n'
     '设计了全新的概率图精炼方法，包含三个关键组件:\n'
     '  (a) 动态阈值 (Dynamic Threshold): MLP 预测每个通道的分位数阈值 α\n'
     '  (b) 局部自适应高斯平滑: MLP 预测平滑参数 σ，对高响应区域做高斯平滑\n'
     '  (c) 可学习温度缩放 (Temperature Scaling): 通道级温度系数调整概率分布\n\n'
     '这种设计使得精炼过程完全可学习、自适应，不需要手动调参。'),

    ('创新3：动态提示编码融合 (Dynamic Prompt Encoding Fusion)',
     '文件: PosFuse.py\n\n'
     '将精炼后的概率图编码为"提示"（Prompt），通过以下方式与主特征融合:\n'
     '  (a) PromptEncoder: 卷积编码 + 3D 随机傅里叶位置编码\n'
     '  (b) SpatialGate: 空间门控机制，生成 (0,1) 注意力图\n'
     '  (c) 残差增强: output = A × (1 + attention)\n\n'
     '与传统 concat/add 融合不同，这种方法让概率图"教会"主网络关注哪里，'
     '而不是简单地叠加信息。'),

    ('创新4：蒙特卡洛采样稳定性 (MC Sampling Stability)',
     '文件: MC_network_backbone.py (mc_refine_prob 方法)\n\n'
     '对编码器中间特征进行多次 softmax 采样（默认5次），取平均后再精炼。\n'
     '蒙特卡洛平均可以:\n'
     '  (a) 降低单次 softmax 的随机波动\n'
     '  (b) 提供更稳定、更可靠的概率估计\n'
     '  (c) 类似于测试时增强（TTA），但集成在训练过程中'),
]
for title, desc in innovations:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_heading('模块依赖关系', level=2)

add_code_block(doc,
    'MC_network_backbone.py (PMGNet 主网络)\n'
    '    ├── pmg_encoder.py (uxnet_conv)          ← 3D ConvNeXt 骨干\n'
    '    ├── mc_refine.py (RefineSegmentation)    ← 概率图精炼\n'
    '    ├── PosFuse.py (ProbPromptFusion)        ← 位置感知融合\n'
    '    └── monai (UnetrBasicBlock, UnetOutBlock) ← 基础模块\n'
    '\n'
    'train_cadic.py\n'
    '    ├── MC_network_backbone.py (PMGNet)      ← 模型\n'
    '    ├── dataset_cadic.py (CTCACSimpleDataset) ← 数据\n'
    '    └── monai (DiceLoss)                     ← 损失函数'
)

doc.add_paragraph()
doc.add_paragraph()

# ========== 结尾 ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

# ========== 保存 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '3D-PMGNet_代码解析文档.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')
